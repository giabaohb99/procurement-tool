"""Thư viện KHỐI dựng file Word chuẩn DEGO — bản port từ skill dego-docx.

Mỗi hàm là một khối trình bày (header, tiêu đề, hộp tóm tắt, thanh mục, bảng...);
tầng trên (export_tool) chỉ compose khối theo dữ liệu, không đụng XML. Tự viết
XML tô nền/viền/ép cột trong python-docx rất dễ sai nên gom hết vào đây.

Gotcha giữ nguyên từ skill gốc:
- Ép bề rộng cột PHẢI qua `set_col_widths` (ghi cả tblGrid lẫn tblW + tắt autofit).
  Chỉ đặt `cell.width` thì Word co giãn theo nội dung -> tràn lề, cắt chữ.
- Màu là RGB 6 hex (Word), KHÔNG phải ARGB 8 hex của xlsx.
- Font mặc định Inter; máy không cài Inter thì Word tự thay — chấp nhận được vì
  file sinh ra để tải về máy người dùng.
"""
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Bảng màu DEGO (RGB 6 hex).
NAVY = "1A4D6B"        # tiêu đề, chữ nhấn
TEAL = "0E7C7B"        # thanh mục, gạch chân header, marker bullet
TEAL_SOFT = "EAF4F4"   # nền nhạt cho h2 / dòng chẵn của bảng
CREAM = "FFF8E1"       # nền hộp TL;DR
ORANGE = "E8A33D"      # viền hộp TL;DR
GRAY = "6B7280"        # phụ đề, ghi chú
WHITE = "FFFFFF"

FONT = "Inter"
USABLE_CM = 17.0       # bề rộng dùng được của trang A4 sau lề (tổng cột bảng <= số này)


# ---------------------------------------------------------------- nền tảng XML

def _set(el, tag: str, **attrs):
    child = OxmlElement(tag)
    for k, v in attrs.items():
        child.set(qn(f"w:{k}"), str(v))
    el.append(child)
    return child


def _p_shading(paragraph, hex_color: str):
    """Tô nền cả đoạn (khác tô nền run)."""
    _set(paragraph._p.get_or_add_pPr(), "w:shd", val="clear", color="auto", fill=hex_color)


def _p_border(paragraph, side: str, hex_color: str, size: int = 8, space: int = 2):
    """Kẻ một cạnh của đoạn (side: top/bottom/left/right)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    _set(pBdr, f"w:{side}", val="single", sz=size, space=space, color=hex_color)


def _cell_shade(cell, hex_color: str):
    _set(cell._tc.get_or_add_tcPr(), "w:shd", val="clear", color="auto", fill=hex_color)


def set_col_widths(table, widths_cm: list[float]):
    """Ép bề rộng cột kiểu cố định — ghi tblGrid + tblW, tắt autofit (gotcha số 1)."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    total_dxa = int(sum(widths_cm) * 567)   # 1 cm = 567 twip
    _set(tblPr, "w:tblW", w=total_dxa, type="dxa")
    _set(tblPr, "w:tblLayout", type="fixed")
    for old in tbl.findall(qn("w:tblGrid")):
        tbl.remove(old)
    grid = OxmlElement("w:tblGrid")
    for w in widths_cm:
        _set(grid, "w:gridCol", w=int(w * 567))
    tbl.insert(list(tbl).index(tblPr) + 1, grid)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_cm):
            cell.width = Cm(w)


def _add_rich(paragraph, text: str, size: float = 10.5, color: str | None = None,
              bold: bool = False):
    """Đổ chuỗi có <strong>/<br> thành runs. Chỉ hỗ trợ đúng 2 thẻ đó — đủ dùng."""
    text = str(text or "")
    for chunk in text.replace("<br/>", "<br>").split("<br>"):
        if paragraph.runs and chunk is not None:
            #  <br> giữa các chunk -> ngắt dòng mềm trong cùng đoạn.
            paragraph.runs[-1].add_break()
        rest = chunk
        while rest:
            if "<strong>" in rest:
                before, _, after = rest.partition("<strong>")
                inner, _, rest = after.partition("</strong>")
                if before:
                    _run(paragraph, before, size, color, bold)
                _run(paragraph, inner, size, NAVY, True)
            else:
                _run(paragraph, rest, size, color, bold)
                rest = ""


def _run(paragraph, text: str, size: float, color: str | None, bold: bool):
    r = paragraph.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return r


# ------------------------------------------------------------------- các khối

def new_document() -> Document:
    """Document A4 dọc, lề chuẩn, font mặc định."""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.0)
    section.top_margin = section.bottom_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    return doc


def header(doc, right_lines: list[str]):
    """Chữ DEGO trái + khối chữ phải, gạch chân teal. (Bản port bỏ logo ảnh —
    môi trường server không kéo pymupdf chỉ để rasterize SVG.)"""
    table = doc.add_table(rows=1, cols=2)
    set_col_widths(table, [7.0, 10.0])
    left = table.rows[0].cells[0].paragraphs[0]
    _run(left, "DEGO HOLDING", 16, NAVY, True)
    right_cell = table.rows[0].cells[1]
    right_p = right_cell.paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for i, line in enumerate(right_lines):
        if i:
            right_p.runs[-1].add_break()
        if i == 0:
            _run(right_p, line, 11, NAVY, True)
        elif i == 1:
            r = _run(right_p, line, 9, GRAY, False)
            r.font.italic = True
        else:
            _run(right_p, line, 9, GRAY, False)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(2)
    rule.paragraph_format.space_after = Pt(8)
    _p_border(rule, "bottom", TEAL, size=16)


def title(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    _run(p, str(text).upper(), 16, NAVY, True)


def subtitle(doc, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = _run(p, text, 10.5, GRAY, False)
    r.font.italic = True


def note(doc, text: str):
    p = doc.add_paragraph()
    r = _run(p, text, 9, GRAY, False)
    r.font.italic = True


def meta_table(doc, pairs: list[tuple[str, str]]):
    """Bảng thông tin chung 2 cột nhãn | giá trị (nhãn tô đậm navy)."""
    if not pairs:
        return
    table = doc.add_table(rows=len(pairs), cols=2)
    table.style = "Table Grid"
    set_col_widths(table, [4.5, USABLE_CM - 4.5])
    for row, (label, value) in zip(table.rows, pairs):
        _cell_shade(row.cells[0], TEAL_SOFT)
        _add_rich(row.cells[0].paragraphs[0], f"<strong>{label}</strong>", 10)
        _add_rich(row.cells[1].paragraphs[0], value, 10)
    spacer(doc, 6)


def tldr_box(doc, heading: str, items: list[str]):
    """Hộp tóm tắt: nền vàng kem, viền cam, bullet teal."""
    head = doc.add_paragraph()
    head.paragraph_format.space_before = Pt(6)
    head.paragraph_format.space_after = Pt(0)
    _p_shading(head, CREAM)
    for side in ("top", "left", "right"):
        _p_border(head, side, ORANGE, size=12)
    _run(head, heading, 11, NAVY, True)
    for i, item in enumerate(items):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0 if i < len(items) - 1 else 0)
        p.paragraph_format.left_indent = Cm(0.3)
        _p_shading(p, CREAM)
        for side in ("left", "right") + (("bottom",) if i == len(items) - 1 else ()):
            _p_border(p, side, ORANGE, size=12)
        _run(p, "»  ", 10.5, TEAL, True)
        _add_rich(p, item, 10.5)
    spacer(doc, 8)


def section_bar(doc, text: str):
    """Thanh mục lớn: nền teal, chữ trắng in hoa."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    _p_shading(p, TEAL)
    _run(p, f"  {str(text).upper()}", 11.5, WHITE, True)


def h2(doc, text: str):
    """Tiêu đề phụ: nền nhạt, viền trái teal đậm."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    _p_shading(p, TEAL_SOFT)
    _p_border(p, "left", TEAL, size=24)
    _run(p, f" {text}", 11, NAVY, True)


def paragraph(doc, text: str):
    p = doc.add_paragraph()
    _add_rich(p, text, 10.5)


def bullet(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    _run(p, "•  ", 10.5, TEAL, True)
    _add_rich(p, text, 10.5)


def data_table(doc, columns: list[str], rows: list[list[str]]):
    """Bảng dữ liệu N cột + cột STT tự đánh; header teal chữ trắng, dòng chẵn tô nền."""
    n = len(columns)
    if n == 0:
        return
    table = doc.add_table(rows=1 + len(rows), cols=n + 1)
    table.style = "Table Grid"
    stt_w = 1.2
    col_w = (USABLE_CM - stt_w) / n
    set_col_widths(table, [stt_w] + [col_w] * n)

    head_cells = table.rows[0].cells
    for cell, text in zip(head_cells, ["STT"] + [str(c) for c in columns]):
        _cell_shade(cell, TEAL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, str(text), 10, WHITE, True)

    for idx, row_values in enumerate(rows, start=1):
        cells = table.rows[idx].cells
        shade = TEAL_SOFT if idx % 2 == 0 else None
        stt_p = cells[0].paragraphs[0]
        stt_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(stt_p, str(idx), 10, NAVY, True)
        if shade:
            _cell_shade(cells[0], shade)
        for cell, value in zip(cells[1:], row_values):
            if shade:
                _cell_shade(cell, shade)
            _add_rich(cell.paragraphs[0], value, 10)
    spacer(doc, 6)


def footer(doc, main_line: str, sub_line: str = ""):
    """Footer lặp mọi trang (đặt trong section footer thật, không phải đoạn cuối)."""
    f = doc.sections[0].footer
    f.is_linked_to_previous = False
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _p_border(p, "top", TEAL, size=8)
    _run(p, main_line, 8.5, GRAY, False)
    if sub_line:
        p.runs[-1].add_break()
        r = _run(p, sub_line, 8, GRAY, False)
        r.font.italic = True


def spacer(doc, pts: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
